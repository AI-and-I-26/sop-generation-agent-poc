"""
custom_sop.py — SOP Generator Entry Point / Test Runner.

INSTRUCTIONS:
    1. Edit the three constants below (TOPIC, INDUSTRY, AUDIENCE).
    2. Run: python test/custom_sop.py
    3. Your SOP will be saved as:
         sop_<topic>.md    — Markdown
         sop_<topic>.docx  — Word document
         sop_<topic>.pdf   — PDF

PREREQUISITES:
    pip install python-docx reportlab strands boto3 pydantic

AWS REQUIREMENTS:
    - AWS credentials configured (env vars or ~/.aws/credentials)
    - Bedrock model access enabled for claude-sonnet-4-6
    - KNOWLEDGE_BASE_ID set to your Bedrock KB (or set in the script below)
"""

# ── Standard library ──────────────────────────────────────────────────────────
import asyncio
import logging
import os
import sys
from pathlib import Path

# ── Logging setup ─────────────────────────────────────────────────────────────
# Set to logging.INFO for normal runs; logging.DEBUG for troubleshooting.
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)

# ── Project root path resolution ──────────────────────────────────────────────
# __file__ is  test/custom_sop.py
# parents[0] is test/
# parents[1] is <project_root>/
# We insert <project_root> into sys.path so all 'src.' imports resolve correctly.
PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

# ── Project import ────────────────────────────────────────────────────────────
from src.graph.sop_workflow import generate_sop   # async entry point


# =============================================================================
# OPTIONAL: Word (.docx) export
# =============================================================================
_DOCX_AVAILABLE = True
try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    _DOCX_AVAILABLE = False
    logging.warning(
        "python-docx not installed — skipping .docx export. "
        "Install with: pip install python-docx"
    )


def _docx_list_para(doc, text: str, numbered: bool = False):
    """Add a bullet or numbered list paragraph using Word's built-in styles."""
    p = doc.add_paragraph(text)
    p.style = "List Number" if numbered else "List Bullet"
    return p


def _docx_code_para(doc, text: str):
    """Add a monospace code-style paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    # Ensure the font is encoded in the XML
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    rPr.append(rFonts)
    return p


def write_markdown_to_docx(doc: "Document", md_text: str, title: str, industry: str, audience: str):
    """
    Lightweight Markdown-to-Word mapper.

    Supports:
      - Headings  (#, ##, ###)
      - Bullet lists  (- or *)
      - Numbered lists  (1. 2. 3.)
      - Fenced code blocks  (``` ... ```)
      - Normal paragraphs
    """
    # Document title and metadata block at the top
    doc.add_heading(title, level=1)
    if industry:
        doc.add_paragraph(f"Industry: {industry}")
    if audience:
        doc.add_paragraph(f"Audience: {audience}")
    doc.add_paragraph("")

    in_code_block = False

    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()

        # Toggle code block on ``` fence markers
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block:
                doc.add_paragraph("")  # spacing after code block
            continue

        if in_code_block:
            _docx_code_para(doc, raw_line)
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)

        # Bullet lists
        elif line.lstrip().startswith("- ") or line.lstrip().startswith("* "):
            _docx_list_para(doc, line.lstrip()[2:].strip(), numbered=False)

        # Numbered lists (e.g. "1. text")
        elif (
            len(line.lstrip()) > 3
            and line.lstrip()[0].isdigit()
            and line.lstrip()[1] == "."
            and line.lstrip()[2] == " "
        ):
            _docx_list_para(doc, line.lstrip()[3:].strip(), numbered=True)

        elif line.strip() == "":
            doc.add_paragraph("")   # blank spacing

        else:
            doc.add_paragraph(raw_line)


# =============================================================================
# OPTIONAL: PDF (.pdf) export via ReportLab
# =============================================================================
_PDF_AVAILABLE = True
try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import (
        ListFlowable, ListItem, Paragraph, Preformatted, SimpleDocTemplate, Spacer
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    _PDF_AVAILABLE = False
    logging.warning(
        "reportlab not installed — skipping .pdf export. "
        "Install with: pip install reportlab"
    )


def _register_unicode_font(font_name: str = "DejaVuSans"):
    """
    Try to register a TrueType Unicode font for broad character support.
    Returns the font name if successful, else None (falls back to Helvetica).
    """
    if not _PDF_AVAILABLE:
        return None
    paths = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",  # Linux
        "/Library/Fonts/DejaVuSans.ttf",                    # macOS (if installed)
        "C:\\Windows\\Fonts\\DejaVuSans.ttf",                # Windows (if installed)
    ]
    for p in paths:
        try:
            if Path(p).is_file():
                pdfmetrics.registerFont(TTFont(font_name, p))
                return font_name
        except Exception:
            continue
    return None


def write_markdown_to_pdf(pdf_path: Path, md_text: str, title: str, industry: str, audience: str):
    """
    Lightweight Markdown-to-PDF renderer using ReportLab Platypus.
    Supports headings, bullets, numbered lists, code blocks, and paragraphs.
    """
    if not _PDF_AVAILABLE:
        raise RuntimeError("ReportLab is not installed.")

    base_font = _register_unicode_font() or "Helvetica"
    mono_font = base_font  # use same font for code (Consolas not in default ReportLab)

    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=LETTER,
        rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("SOP_H1",   parent=styles["Heading1"],  fontName=base_font, fontSize=16, spaceAfter=10))
    styles.add(ParagraphStyle("SOP_H2",   parent=styles["Heading2"],  fontName=base_font, fontSize=13, spaceAfter=8))
    styles.add(ParagraphStyle("SOP_H3",   parent=styles["Heading3"],  fontName=base_font, fontSize=11, spaceAfter=6))
    styles.add(ParagraphStyle("SOP_BODY", parent=styles["BodyText"],  fontName=base_font, leading=14,  spaceAfter=6))
    styles.add(ParagraphStyle("SOP_META", parent=styles["BodyText"],  fontName=base_font, leading=12,  spaceAfter=4))
    styles.add(ParagraphStyle("SOP_CODE", fontName=mono_font, fontSize=9, leading=11, spaceAfter=4))

    flow = []

    # Document header
    if title:    flow.append(Paragraph(title,                    styles["SOP_H1"]))
    if industry: flow.append(Paragraph(f"Industry: {industry}",  styles["SOP_META"]))
    if audience: flow.append(Paragraph(f"Audience: {audience}",  styles["SOP_META"]))
    flow.append(Spacer(1, 10))

    in_code_block = False
    pending_bullets: list = []
    pending_numbers: list = []

    def flush_lists():
        """Flush any accumulated list items to the flow."""
        nonlocal pending_bullets, pending_numbers
        if pending_bullets:
            items = [ListItem(Paragraph(t, styles["SOP_BODY"])) for t in pending_bullets]
            flow.append(ListFlowable(items, bulletType="bullet", start="circle", leftIndent=18))
            pending_bullets = []
            flow.append(Spacer(1, 4))
        if pending_numbers:
            items = [ListItem(Paragraph(t, styles["SOP_BODY"])) for t in pending_numbers]
            flow.append(ListFlowable(items, bulletType="1", leftIndent=18))
            pending_numbers = []
            flow.append(Spacer(1, 4))

    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()

        if line.strip().startswith("```"):
            flush_lists()
            in_code_block = not in_code_block
            if not in_code_block:
                flow.append(Spacer(1, 6))
            continue

        if in_code_block:
            flow.append(Preformatted(raw_line, styles["SOP_CODE"]))
            continue

        if line.startswith("### "):
            flush_lists()
            flow.append(Paragraph(line[4:].strip(), styles["SOP_H3"]))
        elif line.startswith("## "):
            flush_lists()
            flow.append(Paragraph(line[3:].strip(), styles["SOP_H2"]))
        elif line.startswith("# "):
            flush_lists()
            flow.append(Paragraph(line[2:].strip(), styles["SOP_H1"]))
        elif line.lstrip().startswith("- ") or line.lstrip().startswith("* "):
            pending_bullets.append(line.lstrip()[2:].strip())
        elif (
            len(line.lstrip()) > 3
            and line.lstrip()[0].isdigit()
            and line.lstrip()[1] == "."
            and line.lstrip()[2] == " "
        ):
            pending_numbers.append(line.lstrip()[3:].strip())
        elif line.strip() == "":
            flush_lists()
            flow.append(Spacer(1, 6))
        else:
            flush_lists()
            flow.append(Paragraph(raw_line, styles["SOP_BODY"]))

    flush_lists()
    doc.build(flow)


# =============================================================================
# ASYNC MAIN
# =============================================================================

async def main():
    """
    Orchestrates the full SOP generation + export workflow.

    Steps:
      1. Read topic / industry / audience constants.
      2. Call generate_sop() (async Strands graph).
      3. Write the result to .md, .docx, and .pdf files.
      4. Print a summary to stdout.
    """

    # Ensure UTF-8 output on Windows terminals
    try:
        if hasattr(sys.stdout, "reconfigure"):
            sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

    # =========================================================================
    # ✏️  CUSTOMIZE THESE VALUES:
    # =========================================================================
    TOPIC    = "Global Technology Infrastructure Qualification SOP"       # ← e.g. "Employee Onboarding Process"
    INDUSTRY = "Life Science"        # ← e.g. "Healthcare", "Manufacturing"
    AUDIENCE = "IT Qualification Engineers and System Administrators" # ← e.g. "HR Managers and Team Leads"
    # =========================================================================

    print(f"\n{'='*60}")
    print("SOP Generation Starting...")
    print(f"  Topic:    {TOPIC}")
    print(f"  Industry: {INDUSTRY}")
    print(f"  Audience: {AUDIENCE}")
    print(f"{'='*60}\n")

    # ── Generate SOP ──────────────────────────────────────────────────────────
    result = await generate_sop(
        topic=TOPIC,
        industry=INDUSTRY,
        target_audience=AUDIENCE,
    )

    # Use the primary KB-format output; fall back to legacy field
    doc_text = result.formatted_markdown or result.formatted_document

    if not doc_text:
        print("⚠️  WARNING: No formatted document in result. Check logs for errors.")
        print(f"   Status: {result.status}")
        if result.errors:
            print(f"   Errors: {result.errors[-3:]}")
        return

    # ── Safe filename from topic ──────────────────────────────────────────────
    safe_name = (
        TOPIC.lower()
             .replace(" ", "_")
             .replace("—", "")
             .replace(",", "")
             .replace(".", "")
             [:80]      # truncate long names
    )

    # ── Write Markdown ────────────────────────────────────────────────────────
    md_path = Path(f"sop_{safe_name}.md")
    md_path.write_text(doc_text, encoding="utf-8")

    # ── Write Word document ───────────────────────────────────────────────────
    docx_path = None
    if _DOCX_AVAILABLE:
        try:
            docx_path = Path(f"sop_{safe_name}.docx")
            doc = Document()
            write_markdown_to_docx(doc, doc_text, TOPIC, INDUSTRY, AUDIENCE)
            doc.save(str(docx_path))
        except Exception as e:
            logging.exception("Failed to create .docx: %s", e)
            docx_path = None

    # ── Write PDF ─────────────────────────────────────────────────────────────
    pdf_path = None
    if _PDF_AVAILABLE:
        try:
            pdf_path = Path(f"sop_{safe_name}.pdf")
            write_markdown_to_pdf(pdf_path, doc_text, TOPIC, INDUSTRY, AUDIENCE)
        except Exception as e:
            logging.exception("Failed to create .pdf: %s", e)
            pdf_path = None

    # ── Print summary ─────────────────────────────────────────────────────────
    print(f"\n{'='*60}")
    print("✅ SOP Generation Complete!")
    print(f"   Status:        {result.status}")
    print(f"   KB Hits:       {getattr(result, 'kb_hits', 'N/A')}")
    print(f"   Tokens Used:   {getattr(result, 'tokens_used', 'N/A')}")

    if getattr(result, "qa_result", None):
        print(f"   QA Score:      {result.qa_result.score:.1f}/10")
        print(f"   QA Approved:   {result.qa_result.approved}")
        if result.qa_result.issues:
            print(f"   QA Issues:     {len(result.qa_result.issues)}")
            for issue in result.qa_result.issues[:3]:
                print(f"     • {issue}")

    print(f"\n   Markdown:  {md_path}  ({md_path.stat().st_size:,} bytes)")
    print(f"   Word:      {docx_path if docx_path else '(not created)'}")
    print(f"   PDF:       {pdf_path  if pdf_path  else '(not created)'}")
    print(f"{'='*60}\n")


# =============================================================================
# SCRIPT GUARD
# =============================================================================
if __name__ == "__main__":
    asyncio.run(main())