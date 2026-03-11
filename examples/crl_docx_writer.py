# src/utils/crl_docx_writer.py
"""
crl_docx_writer.py — Generates a .docx SOP document with exact CRL-style
per-page header and footer on EVERY page, matching the Charles River
Laboratories standard operating procedure template.

HEADER (every page):
┌─────────────────┬──────────────────────┬──────────────────────────────┐
│  charles river  │ STANDARD OPERATING   │  Doc #: {{document_id}}      │
│    (logo text)  │    PROCEDURE         │  Rev #: {{version}}           │
│                 │                      │  Effective Date: {{eff_date}} │
├─────────────────┴──────────────────────┴──────────────────────────────┤
│  Title: {{title}}                                                       │
├─────────────────────────────────────────────────────────────────────────┤
│         Status CURRENT, Confidential & Proprietary (centred, grey)     │
└─────────────────────────────────────────────────────────────────────────┘

FOOTER (every page):
══════════════════════════════════════════════════════   ← double rule
                                          Page X of Y   ← right-aligned

USAGE
-----
From your pipeline's docx generation step:

    from src.utils.crl_docx_writer import build_crl_docx

    docx_bytes = build_crl_docx(
        title="Global Technology Infrastructure Qualification SOP",
        document_id="GLBL-SOP-00060",
        version="6",
        effective_date="30/Nov/2024",
        markdown_body=state.formatted_markdown,   # or plain text sections
    )
    with open("output.docx", "wb") as f:
        f.write(docx_bytes)

The function also accepts a pre-parsed list of (heading_level, text) tuples
via the `sections` parameter for richer structure control.
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor, Cm, Twips


# ---------------------------------------------------------------------------
# CONSTANTS — matching the CRL template exactly
# ---------------------------------------------------------------------------

_CRL_BRAND_NAME   = "charles river"
_SOP_LABEL        = "STANDARD OPERATING\nPROCEDURE"
_STATUS_LINE      = "Status CURRENT, Confidential & Proprietary"

# Colours (from image analysis)
_GREY_STATUS      = RGBColor(0x80, 0x80, 0x80)   # grey status line text
_BLACK            = RGBColor(0x00, 0x00, 0x00)
_BORDER_COLOR     = "000000"

# Page layout — US Letter, 1-inch margins
_PAGE_WIDTH       = Inches(8.5)
_PAGE_HEIGHT      = Inches(11)
_MARGIN           = Inches(1.0)
_CONTENT_WIDTH_DXA = int((_PAGE_WIDTH - 2 * _MARGIN) / Twips(1))  # ≈ 9360 twips

# Column widths for the 3-column header table (twips)
# Left: ~1.5", Centre: ~4.5", Right: ~3"  — total = content width
_COL_LEFT   = int(Inches(1.5) / Twips(1))
_COL_CENTRE = int(Inches(4.5) / Twips(1))
_COL_RIGHT  = int(Inches(3.0) / Twips(1))


# ---------------------------------------------------------------------------
# XML HELPERS
# ---------------------------------------------------------------------------

def _set_cell_border(cell, **kwargs):
    """Set borders on a table cell. kwargs: top/bottom/left/right = True/False."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    sides = {"top": "w:top", "bottom": "w:bottom", "left": "w:left", "right": "w:right",
             "insideH": "w:insideH", "insideV": "w:insideV"}
    for side, tag in sides.items():
        val = kwargs.get(side, True)
        el = OxmlElement(tag)
        if val:
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    "6")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), _BORDER_COLOR)
        else:
            el.set(qn("w:val"), "none")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _no_space_para(cell):
    """Remove paragraph spacing from all paragraphs in a cell."""
    for p in cell.paragraphs:
        pPr = p._p.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"),  "0")
        pPr.append(spacing)


def _add_run(para, text: str, bold=False, size_pt: int = 9,
             color: RGBColor = _BLACK, italic=False) -> None:
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size_pt)
    run.font.color.rgb = color
    run.font.name  = "Arial"


def _make_table_border(table):
    """Apply single black border to entire table."""
    tbl    = table._tbl
    tblPr  = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), _BORDER_COLOR)
        tblBorders.append(el)
    tblPr.append(tblBorders)
    if tbl.find(qn("w:tblPr")) is None:
        tbl.insert(0, tblPr)


def _page_number_field(para, prefix: str = "Page ", suffix: str = " of "):
    """Insert 'Page X of Y' field runs into a paragraph."""
    _add_run(para, prefix, size_pt=9)

    # PAGE field
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    run = para.add_run(); r = run._r
    r.append(fldBegin); r.append(instrText); r.append(fldEnd)
    run.font.size = Pt(9); run.font.name = "Arial"

    _add_run(para, suffix, size_pt=9)

    # NUMPAGES field
    fldBegin2 = OxmlElement("w:fldChar"); fldBegin2.set(qn("w:fldCharType"), "begin")
    instrText2 = OxmlElement("w:instrText"); instrText2.set(qn("xml:space"), "preserve")
    instrText2.text = " NUMPAGES "
    fldEnd2 = OxmlElement("w:fldChar"); fldEnd2.set(qn("w:fldCharType"), "end")
    run2 = para.add_run(); r2 = run2._r
    r2.append(fldBegin2); r2.append(instrText2); r2.append(fldEnd2)
    run2.font.size = Pt(9); run2.font.name = "Arial"


def _double_rule_para(para):
    """Apply a top double-border to a paragraph (mimics the footer double rule)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "double")
    top.set(qn("w:sz"),    "6")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), _BORDER_COLOR)
    pBdr.append(top)
    pPr.append(pBdr)
    # Remove paragraph spacing
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"),  "60")
    pPr.append(spacing)


# ---------------------------------------------------------------------------
# CRL HEADER BUILDER
# ---------------------------------------------------------------------------

def _build_crl_header(header_obj, title: str, document_id: str,
                      version: str, effective_date: str) -> None:
    """
    Builds the CRL 3-part header table into the given Header object.

    Layout mirrors the uploaded Header.jpg exactly:
    ┌──────────────┬──────────────────────┬────────────────────────────┐
    │ charles river│ STANDARD OPERATING   │ Doc #: GLBL-SOP-00060      │
    │              │    PROCEDURE         │ Rev #: 6                   │
    │              │                      │ Effective Date: 30/Nov/2024│
    ├──────────────┴──────────────────────┴────────────────────────────┤
    │ Title: <title>                                                     │
    ├────────────────────────────────────────────────────────────────────┤
    │          Status CURRENT, Confidential & Proprietary                │
    └────────────────────────────────────────────────────────────────────┘
    """
    # Clear any default empty paragraph
    for p in header_obj.paragraphs:
        p._p.getparent().remove(p._p)

    # ── Row 1: 3-column identity row ────────────────────────────────────
    tbl = header_obj.add_table(rows=1, cols=3, width=Twips(_COL_LEFT + _COL_CENTRE + _COL_RIGHT))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style    = "Table Grid"

    # Set column widths
    for i, width in enumerate([_COL_LEFT, _COL_CENTRE, _COL_RIGHT]):
        for cell in tbl.columns[i].cells:
            cell.width = Twips(width)

    row0 = tbl.rows[0]

    # Cell 0 — "charles river" brand (left-aligned, vertically centred)
    c0 = row0.cells[0]
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p0, _CRL_BRAND_NAME, bold=False, size_pt=12, italic=True)
    _no_space_para(c0)

    # Cell 1 — "STANDARD OPERATING PROCEDURE" (bold, centred)
    c1 = row0.cells[1]
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p1, "STANDARD OPERATING\nPROCEDURE", bold=True, size_pt=11)
    _no_space_para(c1)

    # Cell 2 — Doc meta (right column, smaller font, left-aligned)
    c2 = row0.cells[2]
    c2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p2a = c2.paragraphs[0]
    _add_run(p2a, f"Doc #: {document_id}", size_pt=9)
    p2b = c2.add_paragraph()
    _add_run(p2b, f"Rev #: {version}", size_pt=9)
    p2c = c2.add_paragraph()
    _add_run(p2c, f"Effective Date: {effective_date}", size_pt=9)
    for p in [p2a, p2b, p2c]:
        pPr = p._p.get_or_add_pPr()
        sp  = OxmlElement("w:spacing")
        sp.set(qn("w:before"), "0")
        sp.set(qn("w:after"),  "40")
        pPr.append(sp)

    _make_table_border(tbl)

    # ── Row 2: Title row (merged across all 3 cols) ──────────────────────
    tbl2 = header_obj.add_table(rows=1, cols=1, width=Twips(_COL_LEFT + _COL_CENTRE + _COL_RIGHT))
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl2.style = "Table Grid"
    tbl2.columns[0].cells[0].width = Twips(_COL_LEFT + _COL_CENTRE + _COL_RIGHT)
    title_cell = tbl2.rows[0].cells[0]
    p_title = title_cell.paragraphs[0]
    _add_run(p_title, f"Title: {title}", size_pt=10, bold=False)
    pPr = p_title._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing"); sp.set(qn("w:before"), "40"); sp.set(qn("w:after"), "40")
    pPr.append(sp)
    _make_table_border(tbl2)

    # ── Row 3: Status line (grey, centred) ───────────────────────────────
    p_status = header_obj.add_paragraph()
    p_status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p_status, _STATUS_LINE, size_pt=9, color=_GREY_STATUS, italic=False)
    pPr = p_status._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing"); sp.set(qn("w:before"), "40"); sp.set(qn("w:after"), "0")
    pPr.append(sp)


# ---------------------------------------------------------------------------
# CRL FOOTER BUILDER
# ---------------------------------------------------------------------------

def _build_crl_footer(footer_obj) -> None:
    """
    Builds the CRL footer: double top-rule + right-aligned "Page X of Y".

    Mirrors Footer.jpg exactly:
    ══════════════════════════════════════════  ← double rule
                                  Page 1 of 11  ← right-aligned
    """
    # Remove default empty paragraph
    for p in footer_obj.paragraphs:
        p._p.getparent().remove(p._p)

    # Single paragraph with double top-border and right-aligned page number
    p = footer_obj.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _double_rule_para(p)
    _page_number_field(p, prefix="Page ", suffix=" of ")


# ---------------------------------------------------------------------------
# MARKDOWN → DOCX BODY PARSER
# ---------------------------------------------------------------------------

def _parse_markdown_to_paragraphs(md: str) -> List[Tuple[str, str, int]]:
    """
    Very lightweight Markdown → list of (type, text, level) tuples.
    Types: 'heading', 'table_header', 'table_row', 'table_sep', 'text', 'blank'

    Handles:
    - ## H2  →  ('heading', text, 2)
    - ### H3 →  ('heading', text, 3)
    - | col | → ('table_row', raw_line, 0)
    - |---|  → ('table_sep', '', 0)
    - plain   → ('text', text, 0)
    - italic page-header lines  → skipped (pipeline-injected per-section headers)
    """
    result = []
    for raw in md.splitlines():
        line = raw.rstrip()
        # Skip pipeline-injected per-section page-header italic lines
        if line.startswith("*") and ("Doc #:" in line or "CURRENT" in line):
            continue
        if not line.strip():
            result.append(("blank", "", 0))
            continue
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            result.append(("heading", m.group(2).strip(), level))
            continue
        if line.strip().startswith("|"):
            # Table separator row
            if re.match(r"^\|[\s\-:|]+\|", line.strip()):
                result.append(("table_sep", "", 0))
            else:
                result.append(("table_row", line.strip(), 0))
            continue
        # Strip markdown bold/italic markers for plain text output
        cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        cleaned = re.sub(r"\*(.+?)\*",   r"\1", cleaned)
        cleaned = re.sub(r"> ?",         "",    cleaned)
        result.append(("text", cleaned.strip(), 0))
    return result


def _add_body_content(doc: Document, markdown_body: str) -> None:
    """
    Writes the formatted SOP body into the document.
    Converts Markdown headings → Word headings, tables → Word tables,
    and plain text → Normal paragraphs.
    """
    items = _parse_markdown_to_paragraphs(markdown_body)

    # Collect table rows
    i = 0
    while i < len(items):
        typ, text, level = items[i]

        if typ == "blank":
            i += 1
            continue

        if typ == "heading":
            style = {
                2: "Heading 2",
                3: "Heading 3",
                4: "Heading 4",
            }.get(level, "Heading 2")
            p = doc.add_paragraph(style=style)
            # Strip any residual markdown bold
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p.add_run(clean).bold = (level <= 2)
            i += 1
            continue

        if typ == "table_row":
            # Collect all rows of this table
            table_rows = []
            while i < len(items) and items[i][0] in ("table_row", "table_sep", "blank"):
                if items[i][0] == "table_row":
                    cols = [c.strip() for c in items[i][1].strip("|").split("|")]
                    cols = [re.sub(r"\*\*(.+?)\*\*", r"\1", c) for c in cols]
                    table_rows.append(cols)
                i += 1

            if not table_rows:
                continue

            ncols = max(len(r) for r in table_rows)
            tbl   = doc.add_table(rows=len(table_rows), cols=ncols)
            tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Distribute columns evenly across content width
            col_w = _CONTENT_WIDTH_DXA // ncols
            for ri, row_data in enumerate(table_rows):
                row = tbl.rows[ri]
                for ci in range(ncols):
                    cell = row.cells[ci]
                    cell.width = Twips(col_w)
                    text_val = row_data[ci] if ci < len(row_data) else ""
                    p = cell.paragraphs[0]
                    run = p.add_run(text_val)
                    if ri == 0:      # Header row — bold
                        run.bold = True
                    run.font.size = Pt(9)
                    run.font.name = "Arial"
            continue

        if typ == "text":
            if text.strip():
                p = doc.add_paragraph(style="Normal")
                run = p.add_run(text)
                run.font.size = Pt(10)
                run.font.name = "Arial"
            i += 1
            continue

        i += 1


# ---------------------------------------------------------------------------
# PUBLIC API
# ---------------------------------------------------------------------------

def build_crl_docx(
    title: str,
    document_id: str,
    version: str,
    effective_date: str,
    markdown_body: str,
) -> bytes:
    """
    Build a .docx document with the exact CRL header/footer on every page.

    Parameters
    ----------
    title          : SOP title — appears in header Title row and document body
    document_id    : e.g. "GLBL-SOP-00060"
    version        : e.g. "6" or "1.0"
    effective_date : e.g. "30/Nov/2024"
    markdown_body  : formatted Markdown text from formatter_agent

    Returns
    -------
    bytes — the .docx file content, ready to write to disk or upload to S3.
    """
    doc = Document()

    # ── Page layout ─────────────────────────────────────────────────────
    for section in doc.sections:
        section.page_width    = _PAGE_WIDTH
        section.page_height   = _PAGE_HEIGHT
        section.left_margin   = _MARGIN
        section.right_margin  = _MARGIN
        section.top_margin    = Inches(1.5)   # extra room for the header table
        section.bottom_margin = Inches(1.0)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)

        # Link-to-previous = False so our header/footer is independent
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        _build_crl_header(
            section.header,
            title=title,
            document_id=document_id,
            version=version,
            effective_date=effective_date,
        )
        _build_crl_footer(section.footer)

    # ── Document styles ─────────────────────────────────────────────────
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    for h_style, pts in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        try:
            s = doc.styles[h_style]
            s.font.name = "Arial"
            s.font.size = Pt(pts)
            s.font.bold = True
            s.font.color.rgb = _BLACK
        except KeyError:
            pass

    # ── Body content ────────────────────────────────────────────────────
    _add_body_content(doc, markdown_body)

    # ── Serialise to bytes ──────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# CONVENIENCE: apply to an existing SOPState and write to disk
# ---------------------------------------------------------------------------

def write_crl_docx_from_state(state, output_path: str) -> str:
    """
    Convenience wrapper — pulls all metadata from SOPState and writes .docx.

    Parameters
    ----------
    state       : SOPState instance with formatted_markdown, topic, etc.
    output_path : file path to write the .docx

    Returns
    -------
    output_path (str)
    """
    from datetime import datetime as _dt

    title          = getattr(state, "topic", "Standard Operating Procedure")
    formatted_md   = getattr(state, "formatted_markdown", "") or \
                     getattr(state, "formatted_document", "") or ""
    # Try to pull doc metadata from state if set by formatter
    document_id    = getattr(state, "document_id",    None) or \
                     f"SOP-{_dt.now().strftime('%Y%m%d-%H%M')}"
    version        = getattr(state, "sop_version",    "1.0")
    effective_date = getattr(state, "effective_date", _dt.now().strftime("%d/%b/%Y"))

    docx_bytes = build_crl_docx(
        title=title,
        document_id=document_id,
        version=version,
        effective_date=effective_date,
        markdown_body=formatted_md,
    )
    with open(output_path, "wb") as f:
        f.write(docx_bytes)
    return output_path
